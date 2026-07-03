"""Report generation service for Excel and Word exports."""

import io
import logging

import openpyxl
from docx import Document
from docx.shared import Inches, Pt

from app.schemas.document import DocumentResponse

logger = logging.getLogger(__name__)


class ReportService:
    """Generates Excel and Word files from DocumentResponse data."""

    def generate_excel(self, doc_res: DocumentResponse) -> bytes:
        """Generate an Excel (.xlsx) file containing all issues."""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Issues"

        # Headers
        headers = [
            "Title",
            "Description",
            "Date",
            "AKD",
            "AKD Confidence",
            "Extraction Confidence",
            "Source Sections",
            "Source Pages",
            "Review Flags"
        ]
        ws.append(headers)

        # Style headers
        for col in range(1, len(headers) + 1):
            cell = ws.cell(row=1, column=col)
            cell.font = openpyxl.styles.Font(bold=True)

        # Data rows
        for issue in doc_res.issues:
            row = [
                issue.title,
                issue.description,
                issue.date or "N/A",
                issue.akd or "Unclassified",
                f"{issue.akd_confidence:.2f}" if issue.akd_confidence else "N/A",
                f"{issue.confidence:.2f}",
                ", ".join(issue.sections),
                ", ".join(map(str, issue.source_pages)),
                ", ".join(issue.review_flags) if issue.review_flags else "None",
            ]
            ws.append(row)

        output = io.BytesIO()
        wb.save(output)
        wb.close()
        return output.getvalue()

    def generate_word(self, doc_res: DocumentResponse) -> bytes:
        """Generate a Word (.docx) file containing all issues grouped by AKD."""
        doc = Document()
        
        # Title
        title = doc.add_heading(f"Executive Summary Report - {doc_res.filename}", 0)
        title.alignment = 1 # Center

        # Group issues by AKD
        issues_by_akd = {}
        for issue in doc_res.issues:
            akd = issue.akd or "Unclassified"
            if akd not in issues_by_akd:
                issues_by_akd[akd] = []
            issues_by_akd[akd].append(issue)

        # Sort AKDs alphabetically, but put Unclassified at the end
        sorted_akds = sorted(issues_by_akd.keys())
        if "Unclassified" in sorted_akds:
            sorted_akds.remove("Unclassified")
            sorted_akds.append("Unclassified")

        for akd in sorted_akds:
            doc.add_heading(akd, level=1)
            
            for issue in issues_by_akd[akd]:
                # Issue Title
                p_title = doc.add_paragraph()
                r_title = p_title.add_run(issue.title)
                r_title.bold = True
                r_title.font.size = Pt(12)
                
                # Metadata
                date_str = issue.date or "N/A"
                sections_str = ", ".join(issue.sections)
                p_meta = doc.add_paragraph()
                r_meta = p_meta.add_run(f"Date: {date_str} | Sections: {sections_str}")
                r_meta.italic = True
                r_meta.font.size = Pt(10)
                
                # Warnings if any
                if issue.review_flags:
                    p_warn = doc.add_paragraph()
                    r_warn = p_warn.add_run(f"Warnings: {', '.join(issue.review_flags)}")
                    r_warn.bold = True
                
                # Description
                doc.add_paragraph(issue.description)
                
                doc.add_paragraph() # Spacing

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
